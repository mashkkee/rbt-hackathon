from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import json
import uuid
import re
from datetime import datetime
from typing import Dict, Any, Optional, List
from werkzeug.utils import secure_filename
import psycopg2
from psycopg2.extras import RealDictCursor
import logging
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.text_splitter import CharacterTextSplitter
from langchain_chroma import Chroma
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader
LANGCHAIN_AVAILABLE = True

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. DOCX files won't be supported.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
app.config['UPLOAD_FOLDER'] = 'uploads'

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
if LANGCHAIN_AVAILABLE:
    os.makedirs('chroma', exist_ok=True)

# Constants
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'} if DOCX_AVAILABLE else {'txt', 'pdf'}
DATABASE_URL = os.environ.get("DATABASE_URL")

TRAVEL_KEYWORDS = {
    # Serbian keywords
    "turizam", "putovanje", "hotel", "destinacija", "odmor", "letnovanje", 
    "zimovanje", "avion", "let", "smeštaj", "restoran", "atrakcija", 
    "tura", "izlet", "vodič", "rezervacija", "itinerer", "grad", 
    "zemlja", "plaža", "planina", "muzej", "kultura", "avantura", 
    "krstarenje", "viza", "pasoš", "prtljag", "prevoz", "aerodrom",
    "srbija", "beograd", "novi sad", "niš", "kragujevac", "subotica", 
    "pančevo", "zemun", "vojvodina", "šumadija", "zlatibor", "kopaonik",
    "fruška gora", "đerdap", "ram", "sremski karlovci", "oplenac",
    "studenica", "manasija", "ravanica", "sopoćani", "hilandar",
    "cena", "dinar", "evro", "doručak", "noćenje", "takse", "fakultativni",
    "polazak", "povratak", "transfer", "autobus", "grupa", "putnik",
    # English keywords
    "tourism", "travel", "hotel", "destination", "vacation", "holiday", 
    "flight", "accommodation", "restaurant", "attraction", "tour", 
    "guide", "booking", "itinerary", "city", "country", "beach", 
    "mountain", "museum", "culture", "adventure", "cruise", "resort",
    "visa", "passport", "luggage", "transportation", "airport"
}

# Global variables
llm = None
embeddings = None
vector_store = None
user_sessions = {}

def get_db_connection():
    """Get database connection with error handling"""
    try:
        conn = psycopg2.connect(DATABASE_URL)
        return conn
    except Exception as e:
        logger.error(f"Database connection error: {e}")
        return None

def init_database():
    """Initialize database tables"""
    try:
        conn = get_db_connection()
        if not conn:
            return False
            
        with conn.cursor() as cursor:
            # Create table for travel packages
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS travel_packages (
                    id SERIAL PRIMARY KEY,
                    filename VARCHAR(255) NOT NULL UNIQUE,
                    title VARCHAR(500),
                    description TEXT,
                    destinations JSONB DEFAULT '[]',
                    duration_days INTEGER,
                    duration_nights INTEGER,
                    transport_type VARCHAR(100),
                    dates JSONB DEFAULT '[]',
                    prices JSONB DEFAULT '{}',
                    hotels JSONB DEFAULT '[]',
                    includes JSONB DEFAULT '[]',
                    excludes JSONB DEFAULT '[]',
                    highlights JSONB DEFAULT '[]',
                    contact_info JSONB DEFAULT '{}',
                    raw_content TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Create indexes for faster searches
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_travel_packages_filename ON travel_packages(filename)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_travel_packages_destinations ON travel_packages USING GIN(destinations)")
            
            conn.commit()
            logger.info("Database tables initialized successfully")
            return True
            
    except Exception as e:
        logger.error(f"Database initialization error: {e}")
        return False
    finally:
        if conn:
            conn.close()

def read_docx(file_path: str) -> str:
    """Read DOCX file content"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    
    try:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

def read_file_content(file_path: str) -> str:
    """Read content from various file types"""
    try:
        file_ext = file_path.split('.')[-1].lower()
        
        if file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif file_ext == 'pdf' and LANGCHAIN_AVAILABLE:
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            return "\n".join([page.page_content for page in pages])
        elif file_ext == 'docx':
            return read_docx(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return ""
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        return ""

def extract_structured_data(content: str, filename: str) -> Dict:
    """Extract structured data from travel document using LLM"""
    if not llm or not content.strip():
        return {}
    
    try:
        extraction_prompt = ChatPromptTemplate.from_template("""
        Analiziraj sledeći turistički dokument i izdvoji strukturirane podatke u JSON formatu.
        
        VAŽNO: Izvlači SAMO podatke koji STVARNO postoje u dokumentu. Ne izmišljaj podatke!
        
        Dokument: {content}
        
        Vrati JSON sa sledećim poljima (ostavi prazno ili null ako podaci ne postoje):
        {{
            "title": "Naslov putovanja",
            "description": "Kratak opis putovanja", 
            "destinations": ["lista glavnih destinacija"],
            "duration_days": broj_dana,
            "duration_nights": broj_noćenja,
            "transport_type": "tip prevoza",
            "dates": [
                {{
                    "departure_date": "datum polaska",
                    "return_date": "datum povratka", 
                    "price_regular": cena_regularna,
                    "price_discounted": cena_sa_popustom
                }}
            ],
            "hotels": [
                {{
                    "name": "naziv hotela",
                    "category": "kategorija",
                    "location": "lokacija"
                }}
            ],
            "includes": ["šta je uključeno u cenu"],
            "excludes": ["šta NIJE uključeno u cenu"],
            "highlights": ["glavne atrakcije/aktivnosti"],
            "additional_costs": {{
                "single_room_supplement": iznos,
                "optional_tours": iznos,
                "other": "ostali troškovi"
            }}
        }}
        
        Odgovori SAMO sa JSON-om, bez dodatnog teksta.
        """)
        
        chain = extraction_prompt | llm
        response = chain.invoke({"content": content[:4000]})  # Limit content length
        
        # Parse JSON response
        json_str = response.content.strip()
        
        # Clean up response if it contains markdown formatting
        if json_str.startswith("```json"):
            json_str = json_str.replace("```json", "").replace("```", "").strip()
        elif json_str.startswith("```"):
            json_str = json_str.replace("```", "").strip()
        
        try:
            structured_data = json.loads(json_str)
            return structured_data
        except json.JSONDecodeError as e:
            logger.error(f"JSON parse error for {filename}: {e}")
            logger.error(f"Response was: {json_str}")
            return {}
            
    except Exception as e:
        logger.error(f"Structured data extraction error for {filename}: {e}")
        return {}

def save_to_database(filename: str, structured_data: Dict, raw_content: str) -> bool:
    """Save structured data to database"""
    try:
        conn = get_db_connection()
        if not conn:
            return False
            
        with conn.cursor() as cursor:
            # Check if record already exists
            cursor.execute("SELECT id FROM travel_packages WHERE filename = %s", (filename,))
            existing = cursor.fetchone()
            
            if existing:
                # Update existing record
                cursor.execute("""
                    UPDATE travel_packages SET
                        title = %s,
                        description = %s,
                        destinations = %s,
                        duration_days = %s,
                        duration_nights = %s,
                        transport_type = %s,
                        dates = %s,
                        prices = %s,
                        hotels = %s,
                        includes = %s,
                        excludes = %s,
                        highlights = %s,
                        raw_content = %s,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE filename = %s
                """, (
                    structured_data.get('title'),
                    structured_data.get('description'),
                    json.dumps(structured_data.get('destinations', []), ensure_ascii=False),
                    structured_data.get('duration_days'),
                    structured_data.get('duration_nights'),
                    structured_data.get('transport_type'),
                    json.dumps(structured_data.get('dates', []), ensure_ascii=False),
                    json.dumps(structured_data.get('additional_costs', {}), ensure_ascii=False),
                    json.dumps(structured_data.get('hotels', []), ensure_ascii=False),
                    json.dumps(structured_data.get('includes', []), ensure_ascii=False),
                    json.dumps(structured_data.get('excludes', []), ensure_ascii=False),
                    json.dumps(structured_data.get('highlights', []), ensure_ascii=False),
                    raw_content,
                    filename
                ))
            else:
                # Insert new record
                cursor.execute("""
                    INSERT INTO travel_packages (
                        filename, title, description, destinations, duration_days,
                        duration_nights, transport_type, dates, prices, hotels,
                        includes, excludes, highlights, raw_content
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    filename,
                    structured_data.get('title'),
                    structured_data.get('description'),
                    json.dumps(structured_data.get('destinations', []), ensure_ascii=False),
                    structured_data.get('duration_days'),
                    structured_data.get('duration_nights'),
                    structured_data.get('transport_type'),
                    json.dumps(structured_data.get('dates', []), ensure_ascii=False),
                    json.dumps(structured_data.get('additional_costs', {}), ensure_ascii=False),
                    json.dumps(structured_data.get('hotels', []), ensure_ascii=False),
                    json.dumps(structured_data.get('includes', []), ensure_ascii=False),
                    json.dumps(structured_data.get('excludes', []), ensure_ascii=False),
                    json.dumps(structured_data.get('highlights', []), ensure_ascii=False),
                    raw_content
                ))
            
            conn.commit()
            return True
            
    except Exception as e:
        logger.error(f"Database save error for {filename}: {e}")
        return False
    finally:
        if conn:
            conn.close()

def validate_travel_content(text: str) -> bool:
    """Validate if content is travel/tourism related"""
    if not text.strip():
        return False
    
    text_lower = text.lower()
    keyword_count = sum(1 for keyword in TRAVEL_KEYWORDS if keyword in text_lower)
    return keyword_count >= 3  # Require at least 3 travel keywords

def add_document_to_vector_store(content: str, filename: str) -> bool:
    """Add document to vector store"""
    if not vector_store or not content.strip():
        return False
    
    try:
        # Split text into chunks
        text_splitter = CharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separator="\n"
        )
        chunks = text_splitter.split_text(content)
        
        # Create documents
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": filename, 
                    "chunk_id": i,
                    "upload_date": datetime.now().isoformat()
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # Add to vector store
        vector_store.add_documents(documents)
        logger.info(f"Added {len(documents)} chunks from {filename}")
        return True
        
    except Exception as e:
        logger.error(f"Error adding document to vector store: {e}")
        return False

def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_session_id() -> str:
    """Generate unique session ID"""
    return str(uuid.uuid4())

def get_session_history(session_id: str, max_messages: int = 6) -> str:
    """Get formatted conversation history"""
    if session_id not in user_sessions:
        return ""
    
    messages = user_sessions[session_id]['messages'][-max_messages:]
    history_parts = []
    
    for msg in messages:
        role = "Korisnik" if msg['role'] == 'user' else "TurBot"
        content = msg['content']
        
        # Extract content from JSON if it's bot response
        if msg['role'] == 'assistant':
            try:
                data = json.loads(content)
                content = data.get('content', content)
            except:
                pass
        
        history_parts.append(f"{role}: {content}")
    
    return "\n".join(history_parts)

class TravelBot:
    def __init__(self):
        if not LANGCHAIN_AVAILABLE or not llm:
            self.retrieval_chain = None
            return
            
        self.system_prompt = ChatPromptTemplate.from_messages([
            ("system", """Ti si TurBot, profesionalni turistički asistent za putovanja iz Srbije.

VAŽNO: Tvoj odgovor MORA biti u JSON formatu:
{{"content": "tvoj detaljni odgovor", "reserve": true/false, "gmail": "email@example.com"}}

PRAVILA:
1. content: Daj detaljne informacije koristeći podatke iz konteksta. Samo ukoliko nemas nikakv kontekst.
2. reserve: Postavi na true SAMO ako korisnik jasno želi rezervaciju/booking (npr. "rezerviram", "hoću da idem", "bukiram")
3. gmail: Pronađi email adresu iz konteksta dokumenata. Ako nema email-a, istrazi ga iz svoje baze za tu odredjenu agenciju.

Dostupni kontekst:
{context}

Istorija razgovora:
{history}"""),
            ("human", "{input}")
        ])
        
        if vector_store:
            try:
                self.qa_chain = create_stuff_documents_chain(llm, self.system_prompt)
                self.retrieval_chain = create_retrieval_chain(
                    retriever=vector_store.as_retriever(search_kwargs={"k": 5}),
                    combine_docs_chain=self.qa_chain
                )
            except Exception as e:
                logger.error(f"Error creating retrieval chain: {e}")
                self.retrieval_chain = None
        else:
            self.retrieval_chain = None

    def process_message(self, message: str, session_history: str = "") -> Dict[str, Any]:
        """Process user message using RAG or fallback response"""
        if not self.retrieval_chain:
            return self._fallback_response(message)

        try:
            # Use RAG to get response with context
            result = self.retrieval_chain.invoke({
                "input": message,
                "history": session_history
            })
            
            response_text = result.get("answer", "")
            # Try to parse JSON response
            try:
                # Extract JSON from response
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                
                if start_idx != -1 and end_idx > start_idx:
                    json_str = response_text[start_idx:end_idx]
                    response_data = json.loads(json_str)
                else:
                    raise ValueError("No JSON found")
                    
            except (json.JSONDecodeError, ValueError):
                # Fallback if JSON parsing fails
                response_data = {
                    "content": response_text or "Izvinjavam se, nemam odgovor na vaše pitanje.",
                    "reserve": False,
                    "gmail": ""
                }
            
            # Ensure all required fields exist
            response_data.setdefault("content", "")
            response_data.setdefault("reserve", False)
            response_data.setdefault("gmail", "")
            
            return response_data
            
        except Exception as e:
            logger.error(f"Error processing message: {e}")
            return {
                "content": "Došlo je do greške. Molim pokušajte ponovo.",
                "reserve": False,
                "gmail": ""
            }
    
    def _fallback_response(self, message: str) -> Dict[str, Any]:
        """Provide fallback response when LLM is not available"""
        message_lower = message.lower()
        
        # Check for reservation intent
        reserve_keywords = ["rezerviram", "hoću da idem", "bukiram", "rezervacija", "booking"]
        reserve_intent = any(keyword in message_lower for keyword in reserve_keywords)
        
        if any(keyword in message_lower for keyword in ["zdravo", "pozdrav", "hello", "hi"]):
            content = "Zdravo! Ja sam TurBot, vaš turistički asistent. Mogu vam pomoći sa informacijama o putovanjima i turističkim aranžmanima."
        elif any(keyword in message_lower for keyword in ["hvala", "thanks", "thank you"]):
            content = "Nema na čemu! Tu sam da pomognem sa vašim turističkim potrebama."
        else:
            content = "Trenutno nemam pristup bazi turističkih podataka. Molim vas otpremite turistička dokumenta ili kontaktirajte direktno turističku agenciju za detaljne informacije."
        
        return {
            "content": content,
            "reserve": reserve_intent,
            "gmail": ""
        }

# Initialize components
def init_components():
    """Initialize LLM, embeddings, and vector store"""
    global llm, embeddings, vector_store
    
    if not LANGCHAIN_AVAILABLE:
        logger.warning("LangChain not available - limited functionality")
        return
    
    # Initialize LLM
    api_key = os.environ.get("OPENAI_API_KEY")
    if api_key:
        try:
            llm = ChatOpenAI(
                model_name="gpt-4o-mini",
                temperature=0.3,
                max_tokens=1500,
                base_url="https://models.inference.ai.azure.com",
                api_key=api_key
            )
            logger.info("✓ LLM initialized successfully")
        except Exception as e:
            logger.error(f"✗ Failed to initialize LLM: {e}")
    else:
        logger.warning("OPENAI_API_KEY not set - LLM functionality disabled")

    # Initialize embeddings
    if api_key:
        try:
            embeddings = OpenAIEmbeddings(
                base_url="https://models.inference.ai.azure.com",
                api_key=api_key,
                model="text-embedding-3-large"
            )
            logger.info("✓ Embeddings initialized successfully")
        except Exception as e:
            logger.error(f"✗ Failed to initialize embeddings: {e}")

    # Initialize vector store
    if embeddings:
        try:
            vector_store = Chroma(
                collection_name="travel_docs",
                embedding_function=embeddings,
                persist_directory="chroma"
            )
            logger.info("✓ Vector store initialized successfully")

            # Fetch all stored documents (with metadata)
            results = vector_store.similarity_search("test", k=5)

        except Exception as e:
            logger.error(f"✗ Failed to initialize vector store: {e}")

# Routes
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '1.0.0',
        'langchain_available': LANGCHAIN_AVAILABLE,
        'docx_available': DOCX_AVAILABLE,
        'llm_available': llm is not None,
        'vector_store_available': vector_store is not None,
        'database_available': get_db_connection() is not None
    })

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Upload and process travel documents"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            extensions_str = ', '.join(ALLOWED_EXTENSIONS)
            return jsonify({'error': f'File type not supported. Allowed: {extensions_str}'}), 400
        
        # Save file with timestamp
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{timestamp}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Read and process content
        content = read_file_content(file_path)
        if not content:
            os.remove(file_path)
            return jsonify({'error': 'Could not read file content'}), 400
        
        # Validate travel content
        if not validate_travel_content(content):
            os.remove(file_path)
            return jsonify({'error': 'Document does not appear to be travel/tourism related'}), 400
        
        # Extract structured data if LLM is available
        structured_data = {}
        if llm:
            structured_data = extract_structured_data(content, filename)
        
        # Save to database
        db_saved = save_to_database(filename, structured_data, content)
        
        # Add to vector store
        vector_success = False
        if vector_store:
            vector_success = add_document_to_vector_store(content, filename)
        
        return jsonify({
            'message': f'File {filename} uploaded and processed successfully',
            'filename': filename,
            'content_length': len(content),
            'structured_data': structured_data,
            'saved_to_database': db_saved,
            'added_to_vector_store': vector_success,
            'upload_date': datetime.now().isoformat()
        })
            
    except Exception as e:
        logger.error(f"Upload error: {e}")
        return jsonify({'error': 'Upload failed'}), 500


@app.route('/api/upload-multiple', methods=['POST'])
def upload_multiple_files():
    """Upload and process multiple travel documents"""
    try:
        files = request.files.getlist('files')
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        results = []
        for file in files:
            if file.filename == '':
                results.append({'filename': '', 'error': 'No file selected'})
                continue

            if not allowed_file(file.filename):
                extensions_str = ', '.join(ALLOWED_EXTENSIONS)
                results.append({'filename': file.filename, 'error': f'File type not supported. Allowed: {extensions_str}'})
                continue

            # Save file with timestamp
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            full_filename = f"{timestamp}_{filename}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], full_filename)
            file.save(file_path)

            # Read and process content
            content = read_file_content(file_path)
            if not content:
                os.remove(file_path)
                results.append({'filename': full_filename, 'error': 'Could not read file content'})
                continue

            # Validate travel content
            if not validate_travel_content(content):
                os.remove(file_path)
                results.append({'filename': full_filename, 'error': 'Document not travel-related'})
                continue

            # Extract structured data
            structured_data = extract_structured_data(content, full_filename) if llm else {}

            # Save to DB and vector store
            db_saved = save_to_database(full_filename, structured_data, content)
            vector_success = add_document_to_vector_store(content, full_filename) if vector_store else False

            results.append({
                'filename': full_filename,
                'content_length': len(content),
                'structured_data': structured_data,
                'saved_to_database': db_saved,
                'added_to_vector_store': vector_success,
                'upload_date': datetime.now().isoformat()
            })

        return jsonify({'results': results})
    
    except Exception as e:
        logger.error(f"Multiple upload error: {e}")
        return jsonify({'error': 'Multiple file upload failed'}), 500


@app.route('/api/chat', methods=['POST'])
def chat():
    """Main chat endpoint"""
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return jsonify({'error': 'Message is required'}), 400

        user_message = data['message'].strip()
        if not user_message:
            return jsonify({'error': 'Message cannot be empty'}), 400

        # Get or create session
        session_id = data.get('session_id') or generate_session_id()
        if session_id not in user_sessions:
            user_sessions[session_id] = {
                'created_at': datetime.now(),
                'messages': []
            }

        # Add user message to session
        user_sessions[session_id]['messages'].append({
            'role': 'user',
            'content': user_message,
            'timestamp': datetime.now().isoformat()
        })

        # Get conversation history
        history = get_session_history(session_id)

        # Process message
        start_time = datetime.now()
        response_data = travel_bot.process_message(user_message, history)
        processing_time = (datetime.now() - start_time).total_seconds()

        # Add bot response to session
        user_sessions[session_id]['messages'].append({
            'role': 'assistant',
            'content': json.dumps(response_data, ensure_ascii=False),
            'timestamp': datetime.now().isoformat()
        })

        return jsonify({
            'response': response_data,
            'session_id': session_id,
            'processing_time': processing_time,
            'timestamp': datetime.now().isoformat()
        })

    except Exception as e:
        logger.error(f"Chat error: {e}")
        return jsonify({
            'response': {
                "content": "Došlo je do greške u sistemu. Molim pokušajte ponovo.",
                "reserve": False,
                "gmail": ""
            },
            'error': str(e)
        }), 500

@app.route('/api/travel-packages', methods=['GET'])
def get_travel_packages():
    """Get all travel packages from database"""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500

        with conn.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, filename, title, description, destinations, 
                       duration_days, duration_nights, transport_type,
                       dates, prices, hotels, includes, excludes, highlights,
                       created_at, updated_at
                FROM travel_packages 
                ORDER BY created_at DESC
            """)
            packages = cursor.fetchall()

            result = []
            for package in packages:
                package_dict = dict(package)
                # Ensure JSON fields are properly handled
                for field in ['destinations', 'dates', 'prices', 'hotels', 'includes', 'excludes', 'highlights']:
                    if package_dict[field] is None:
                        package_dict[field] = [] if field != 'prices' else {}
                result.append(package_dict)

            return jsonify({
                'packages': result,
                'total': len(result)
            })

    except Exception as e:
        logger.error(f"Get travel packages error: {e}")
        return jsonify({'error': 'Failed to retrieve travel packages'}), 500
    finally:
        if conn:
            conn.close()

# Error handlers
@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 50MB.'}), 413

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {e}")
    return jsonify({'error': 'Internal server error'}), 500

# Initialize everything
init_database()
init_components()
travel_bot = TravelBot()

if __name__ == '__main__':
    logger.info("Starting TurBot Flask API Server...")
    logger.info(f"Upload folder: {app.config['UPLOAD_FOLDER']}")
    
    # Check environment variables
    if not os.environ.get("OPENAI_API_KEY"):
        logger.warning("OPENAI_API_KEY not set - LLM functionality will be limited")
    
    # Check component availability
    logger.info(f"LangChain available: {LANGCHAIN_AVAILABLE}")
    logger.info(f"DOCX support: {DOCX_AVAILABLE}")
    logger.info(f"LLM available: {llm is not None}")
    logger.info(f"Embeddings available: {embeddings is not None}")
    logger.info(f"Vector store available: {vector_store is not None}")
    
    app.run(
        host='0.0.0.0',
        port=8000,
        debug=os.environ.get('FLASK_ENV') == 'development'
    )